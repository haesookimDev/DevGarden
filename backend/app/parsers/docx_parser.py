from docx import Document

from app.parsers.base import DocumentParser, ParsedDocument


class DOCXParser(DocumentParser):
    def parse(self, file_data: bytes) -> ParsedDocument:
        import io

        doc = Document(io.BytesIO(file_data))
        text_parts = []
        images = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            text_parts.append("\n".join(table_text))

        # Extract images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                images.append(rel.target_part.blob)

        return ParsedDocument(
            text="\n\n".join(text_parts),
            images=images,
            metadata={},
            page_count=len(doc.paragraphs) // 30 or 1,  # Rough estimate
        )
